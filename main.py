from flask import Flask
from flask import render_template
from flask import request
from flask import flash, send_file

from flask_bootstrap import Bootstrap
app = Flask(__name__)
app.secret_key = b'_5#1ay21L"F4Q824z\n\xec]/'

from docx import Document
import re

Bootstrap(app)

from wtforms import Form, BooleanField, StringField, PasswordField, validators

class SurveyForm(Form):
    title = StringField('Name of Project', [validators.Length(min=4)])
    address = StringField('Street Address', [validators.Length(min=4)])
    zipcode = StringField('90006', [validators.Length(min=5, max=5)])
    #email = StringField('Email Address', [validators.Length(min=6, max=35)])
    #password = PasswordField('New Password', [
    #    validators.DataRequired(),
    #    validators.EqualTo('confirm', message='Passwords must match')
    #])
    #confirm = PasswordField('Repeat Password')
    #accept_tos = BooleanField('I accept the TOS', [validators.DataRequired()])

def docx_replace_regex(doc_obj, regex , replace):
    for p in doc_obj.paragraphs:
        if regex.search(p.text):
            inline = p.runs
            # Loop added to work with runs (strings with same style)
            for i in range(len(inline)):
                if regex.search(inline[i].text):
                    text = regex.sub(replace, inline[i].text)
                    inline[i].text = text

    for table in doc_obj.tables:
        for row in table.rows:
            for cell in row.cells:
                docx_replace_regex(cell, regex , replace)
    
    if hasattr(doc_obj, 'sections'):
        for section in doc_obj.sections:
            print(section.footer.paragraphs[0].text)
            docx_replace_regex(section.footer, regex, replace)

class Report():
    def __init__(self, form):
        self.form = form
        self.template = Document('template.docx')
        print(self.template.sections[0].header)
        print(self.template.sections[0].footer)
    
    def generate(self):
        for field in self.form:
            docx_replace_regex(self.template, re.compile(field.label.text, re.I), field.data)
        self.filename = 'report.docx'
        self.template.save(self.filename)


@app.route('/', methods=['GET', 'POST'])
def main():
    form = SurveyForm(request.form)
    if request.method == 'POST' and form.validate():
        report = Report(form)
        report.generate()
        flash('Making survey report...')
        return send_file(filename_or_fp=report.filename, mimetype='application/html', as_attachment=True, attachment_filename=report.filename)
    return render_template('main.html', form=form)
